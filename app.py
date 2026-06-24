import io
import json
import os
import re
import subprocess
import tempfile
from pathlib import Path
from typing import List, Tuple

import mammoth
import streamlit as st
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from markdown import markdown
from openai import OpenAI

APP_NAME = "Dadi Translator"
LOGO_PATH = Path(__file__).parent / "assets" / "dadi-coach-logo.png"
DEFAULT_MODEL = os.getenv("OPENAI_MODEL", "gpt-4.1-mini")
MAX_CHARS_PER_CHUNK = 6500

st.set_page_config(
    page_title=APP_NAME,
    page_icon="🌐",
    layout="wide",
    initial_sidebar_state="expanded",
)

CUSTOM_CSS = """
<style>
    .main .block-container { padding-top: 2rem; padding-bottom: 3rem; max-width: 1180px; }
    .dadi-hero {
        border: 1px solid #dbe8d0;
        background: linear-gradient(135deg, #ffffff 0%, #f7fff0 55%, #fff8e6 100%);
        border-radius: 22px;
        padding: 28px 30px;
        margin-bottom: 22px;
    }
    .dadi-title { font-size: 38px; font-weight: 800; color: #255c17; margin: 0; }
    .dadi-subtitle { font-size: 17px; color: #4d5b44; margin-top: 8px; line-height: 1.5; }
    .small-note { color: #697160; font-size: 13px; }
    .stButton > button {
        border-radius: 12px;
        font-weight: 700;
        border: 1px solid #74bd20;
    }
    .stDownloadButton > button {
        border-radius: 12px;
        font-weight: 700;
        background-color: #76bd1d;
        color: white;
        border: 0;
    }
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)


def get_secret_value(key: str, default: str = "") -> str:
    try:
        return st.secrets.get(key, default)
    except Exception:
        return default


def sanitize_filename(name: str) -> str:
    name = re.sub(r"[\\/:*?\"<>|]", "", name).strip()
    name = re.sub(r"\s+", " ", name)
    return name[:120] or "Dadi Translated Document"


def clean_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_title_from_text(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return "Dadi Translated Document"
    first = lines[0]
    first = re.sub(r"^[#\-\*\d\.、\s]+", "", first).strip()
    return first[:80] or "Dadi Translated Document"


def extract_docx_text(uploaded_file) -> str:
    uploaded_file.seek(0)
    result = mammoth.extract_raw_text(uploaded_file)
    return clean_text(result.value)


def extract_txt_text(uploaded_file) -> str:
    raw = uploaded_file.getvalue()
    for encoding in ["utf-8-sig", "utf-8", "gb18030", "big5", "latin-1"]:
        try:
            return clean_text(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    return clean_text(raw.decode("utf-8", errors="ignore"))


def extract_md_text(uploaded_file) -> str:
    md_text = extract_txt_text(uploaded_file)
    html = markdown(md_text)
    soup = BeautifulSoup(html, "html.parser")
    return clean_text(soup.get_text("\n"))


def extract_doc_text(uploaded_file) -> Tuple[str, str]:
    """Legacy .doc files are intentionally not processed in the cloud-safe build.

    Old Microsoft Word .doc files require external system tools that can fail during
    Streamlit Cloud installation. For reliable deployment, users should convert .doc
    files to .docx before upload.
    """
    return "", "Legacy .doc files are not supported in the cloud-safe version. Please open the file in Microsoft Word or WPS, save it as .docx, then upload it again."


def extract_uploaded_text(uploaded_file) -> Tuple[str, str]:
    ext = Path(uploaded_file.name).suffix.lower()
    try:
        if ext == ".docx":
            return extract_docx_text(uploaded_file), ""
        if ext == ".txt":
            return extract_txt_text(uploaded_file), ""
        if ext == ".md":
            return extract_md_text(uploaded_file), ""
        if ext == ".doc":
            return extract_doc_text(uploaded_file)
        return "", "Unsupported file type. Please upload DOCX, DOC, TXT, or MD."
    except Exception as exc:
        return "", f"File extraction failed: {exc}"


def split_text_into_chunks(text: str, max_chars: int = MAX_CHARS_PER_CHUNK) -> List[str]:
    paragraphs = text.split("\n")
    chunks: List[str] = []
    current: List[str] = []
    current_len = 0
    for paragraph in paragraphs:
        extra_len = len(paragraph) + 1
        if current and current_len + extra_len > max_chars:
            chunks.append("\n".join(current).strip())
            current = []
            current_len = 0
        current.append(paragraph)
        current_len += extra_len
    if current:
        chunks.append("\n".join(current).strip())
    return [chunk for chunk in chunks if chunk]


def openai_client(api_key: str) -> OpenAI:
    return OpenAI(api_key=api_key)


def translate_title(client: OpenAI, model: str, chinese_title: str) -> str:
    prompt = (
        "Translate this Simplified Chinese document title into concise professional English. "
        "Return only the English title, without quotation marks.\n\n"
        f"Chinese title:\n{chinese_title}"
    )
    response = client.chat.completions.create(
        model=model,
        temperature=0.2,
        messages=[
            {"role": "system", "content": "You are a professional Chinese-to-English business translator."},
            {"role": "user", "content": prompt},
        ],
    )
    return response.choices[0].message.content.strip() or "Dadi Translated Document"


def translate_chunk(client: OpenAI, model: str, chunk: str, index: int, total: int) -> str:
    system_prompt = (
        "You are the translation engine for Dadi Translator. Translate Simplified Chinese into professional English. "
        "Preserve the original meaning, section order, numbering, bullets, paragraph breaks, and business tone. "
        "Do not summarize. Do not add new information. Do not create bilingual columns. "
        "Return only the English translation of the provided content."
    )
    user_prompt = (
        f"Translate this document segment into professional English. Segment {index} of {total}.\n\n"
        "Formatting rules:\n"
        "1. Preserve headings, numbering, bullet points, and paragraph flow.\n"
        "2. Use clear professional business English.\n"
        "3. Keep terms such as Dadi, Coach, ABCmouse, Clark, and ecosystem campus consistent.\n"
        "4. Return only the translated English text.\n\n"
        f"Chinese text:\n{chunk}"
    )
    response = client.chat.completions.create(
        model=model,
        temperature=0.2,
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    )
    return response.choices[0].message.content.strip()


def translate_document(api_key: str, model: str, chinese_text: str) -> Tuple[str, str]:
    client = openai_client(api_key)
    chinese_title = extract_title_from_text(chinese_text)
    english_title = translate_title(client, model, chinese_title)
    chunks = split_text_into_chunks(chinese_text)
    translated_chunks = []
    progress = st.progress(0)
    status = st.empty()
    for i, chunk in enumerate(chunks, start=1):
        status.info(f"Translating segment {i} of {len(chunks)}...")
        translated_chunks.append(translate_chunk(client, model, chunk, i, len(chunks)))
        progress.progress(i / len(chunks))
    status.success("Translation completed.")
    return english_title, clean_text("\n\n".join(translated_chunks))


def is_heading(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return False
    if len(stripped) <= 90 and re.match(r"^(Chapter|Section|Part|[0-9]+[\.|\-|—]|[IVXLC]+\.|[A-Z][A-Za-z ]+:)", stripped):
        return True
    if len(stripped) <= 70 and stripped.endswith(":"):
        return True
    return False


def create_word_document(title: str, body: str) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    header = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if LOGO_PATH.exists():
        try:
            header_para.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
        except Exception:
            header_para.add_run(APP_NAME)
    else:
        header_para.add_run(APP_NAME)

    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Dadi Translator | Generated English Translation")
    footer_run.font.size = Pt(9)

    title_para = document.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(18)

    subtitle_para = document.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Professional English Translation")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10)

    document.add_paragraph()

    for raw_line in body.splitlines():
        line = raw_line.strip()
        if not line:
            document.add_paragraph()
            continue
        if re.match(r"^[•\-·]\s+", line):
            para = document.add_paragraph(style="List Bullet")
            text = re.sub(r"^[•\-·]\s+", "", line)
            run = para.add_run(text)
        elif re.match(r"^\d+[\.)]\s+", line):
            para = document.add_paragraph(style="List Number")
            text = re.sub(r"^\d+[\.)]\s+", "", line)
            run = para.add_run(text)
        else:
            para = document.add_paragraph()
            run = para.add_run(line)
            if is_heading(line):
                run.bold = True
                run.font.size = Pt(13)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        run.font.name = "Calibri"
        if run.font.size is None:
            run.font.size = Pt(11)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def render_header():
    col1, col2 = st.columns([1, 4])
    with col1:
        if LOGO_PATH.exists():
            st.image(str(LOGO_PATH), use_container_width=True)
    with col2:
        st.markdown(
            """
            <div class='dadi-hero'>
                <p class='dadi-title'>Dadi Translator</p>
                <p class='dadi-subtitle'>Upload a Simplified Chinese document, translate it into professional English, preview the result, and download a formatted Word document.</p>
            </div>
            """,
            unsafe_allow_html=True,
        )


def main():
    render_header()

    with st.sidebar:
        st.subheader("Settings")
        secret_key = get_secret_value("OPENAI_API_KEY", "")
        api_key_input = st.text_input(
            "OpenAI API Key",
            value="" if secret_key else "",
            type="password",
            help="Add this in Streamlit Secrets for deployment. You may also enter it here for local testing.",
        )
        api_key = secret_key or api_key_input
        model = get_secret_value("OPENAI_MODEL", DEFAULT_MODEL)
        model = st.text_input("Model", value=model)
        st.markdown("<p class='small-note'>Recommended cloud-safe version: DOCX, TXT, MD. For legacy DOC, please convert to DOCX first.</p>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader(
        "Upload Simplified Chinese document",
        type=["docx", "doc", "txt", "md"],
        help="Upload a Chinese document. The app will extract the text and translate it into English.",
    )

    if uploaded_file:
        text, warning = extract_uploaded_text(uploaded_file)
        if warning:
            st.warning(warning)
        if text:
            st.success(f"Extracted {len(text):,} characters from {uploaded_file.name}.")
            with st.expander("Preview extracted Chinese text", expanded=False):
                st.text_area("Chinese source text", text, height=260)

            if "source_text" not in st.session_state or st.session_state.get("file_name") != uploaded_file.name:
                st.session_state.source_text = text
                st.session_state.file_name = uploaded_file.name
                st.session_state.translated_title = ""
                st.session_state.translated_body = ""

            translate_clicked = st.button("Translate to English", type="primary", use_container_width=True)
            if translate_clicked:
                if not api_key:
                    st.error("Please add your OpenAI API key in the sidebar or Streamlit Secrets.")
                else:
                    try:
                        title, translated = translate_document(api_key, model, st.session_state.source_text)
                        st.session_state.translated_title = title
                        st.session_state.translated_body = translated
                    except Exception as exc:
                        st.error(f"Translation failed: {exc}")

    if st.session_state.get("translated_body"):
        st.divider()
        st.subheader("Review and Edit English Translation")
        edited_title = st.text_input("Document Title", value=st.session_state.translated_title)
        edited_body = st.text_area("English Translation", value=st.session_state.translated_body, height=520)

        docx_bytes = create_word_document(edited_title, edited_body)
        file_name = sanitize_filename(edited_title) + ".docx"
        st.download_button(
            label="Download Word Document",
            data=docx_bytes,
            file_name=file_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
        st.caption(f"Generated file name: {file_name}")


if __name__ == "__main__":
    main()
